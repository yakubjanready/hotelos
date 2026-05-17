#!/usr/bin/env python3
"""
HotelOS yakuniy hisobotni Word formatida yaratuvchi skript.

Talablar:
- A4, Times New Roman 12pt, 1.5 oraliq
- Chegaralar: Chap 3sm, O'ng 1.5sm, Yuqori 2sm, Pastki 2sm
- Tekis hizalash, 1.25sm birinchi qator chekinishi
- Mundarijadan boshlab sahifa raqamlari pastki markazda
- Harvard formatida 8+ manba

Ishga tushirish:
    python3 scripts/generate_report.py
Chiqish:
    docs/HotelOS_Hisobot.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor

DOCS = Path(__file__).resolve().parent.parent / "docs"


# ============================================================
#  STILLAR VA YORDAMCHI FUNKSIYALAR
# ============================================================

def setup_document_styles(doc: Document) -> None:
    """Times New Roman 12pt, 1.5 oraliq sukut bo'yicha."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")

    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    # Heading styles
    for lvl, size, bold in [(1, 16, True), (2, 14, True), (3, 13, True)]:
        h = doc.styles[f"Heading {lvl}"]
        h.font.name = "Times New Roman"
        h.font.size = Pt(size)
        h.font.bold = bold
        h.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

    # Page setup: A4, custom margins
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)


def add_footer_page_numbers(doc: Document) -> None:
    """Sahifa raqamlari pastki markazda."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_para(doc: Document, text: str, *, justify: bool = True, indent: bool = True,
             bold: bool = False, italic: bool = False, size: int | None = None):
    """Asosiy matn paragrafini qo'shadi.

    indent=True birinchi qatorni 1.25sm chekinadi (brief talabi).
    """
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    """Sarlavha qo'shadi (chekinishsiz)."""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.first_line_indent = Cm(0)
    return h


def add_code(doc: Document, code: str, language: str = "Python"):
    """Monoslab shriftda kod parchasini qo'shadi (kulrang fon bilan)."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.5)
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    # Background shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F1F5F9")
    pPr.append(shd)

    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    # rFonts override
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    rFonts.set(qn("w:cs"), "Consolas")
    return p


def add_image(doc: Document, path: Path, caption: str, width_cm: float = 16):
    """Rasmni qo'shadi va pastida raqamlangan sarlavha."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(str(path), width=Cm(width_cm))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(11)
    return cap


def add_placeholder_screenshot(doc: Document, caption: str, instruction: str):
    """Skrinshot uchun joy egasi (rasm tayyor bo'lguncha) — siz lokal mashinada olib qo'yasiz."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FEF3C7")
    pPr.append(shd)
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    border = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "8"); b.set(qn("w:color"), "B45309")
        border.append(b)
    pPr.append(border)
    run = p.add_run(f"[ SKRINSHOT KUTILMOQDA ]\n{instruction}")
    run.bold = True
    run.font.size = Pt(11)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              col_widths_cm: list[float] | None = None):
    """Yuqorida sarlavhasi va chegaralari bo'lgan oddiy jadval."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    if col_widths_cm:
        for col, w in zip(table.columns, col_widths_cm):
            for cell in col.cells:
                cell.width = Cm(w)
    # Headers
    for i, head in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(head)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    # Rows
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
    return table


def add_table_caption(doc: Document, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(11)


def page_break(doc: Document):
    doc.add_page_break()


# ============================================================
#  MUQOVA VA MUNDARIJA
# ============================================================

def add_cover_page(doc: Document):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BTEC Pearson Raqamli Texnologiyalar")
    run.bold = True; run.font.size = Pt(14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("4-Modul: Dasturlash (H/618/7388)")
    run.bold = True; run.font.size = Pt(14)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("HOTELOS")
    run.bold = True; run.font.size = Pt(36); run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Real Vaqtli Mehmonxona Boshqaruv Tizimi")
    run.italic = True; run.font.size = Pt(18)

    for _ in range(4):
        doc.add_paragraph()

    info = [
        ("Talaba ismi:", "Yoqubjon"),
        ("Talaba ID:", "[ID ni kiriting]"),
        ("Modul nomi:", "4-Modul: Dasturlash"),
        ("Vazifa nomi:", "HotelOS: Real Vaqtli Mehmonxona Boshqaruv Tizimini Qurish"),
        ("Baholovchi:", "[Baholovchi ismini kiriting]"),
        ("Topshirish sanasi:", "17 May 2026"),
    ]
    for label, value in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        r1 = p.add_run(label + " "); r1.bold = True; r1.font.size = Pt(12)
        r2 = p.add_run(value); r2.font.size = Pt(12)


def add_toc(doc: Document):
    p = doc.add_paragraph()
    run = p.add_run("Mundarija")
    run.bold = True
    run.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)

    doc.add_paragraph()
    items = [
        ("1-Vazifa — Algoritmlar va Kod Jarayoni (LO1)", "4"),
        ("    1.1 Xona Tayinlash Algoritmi", "4"),
        ("    1.2 Qo'shimcha Algoritmlar (Hisob-kitob, Priority Queue)", "7"),
        ("    1.3 Koddan Bajarilishgacha — Python Ish Vaqti", "9"),
        ("    1.4 Texnologiya Stekini Asoslash", "10"),
        ("2-Vazifa — HotelOSda Dasturlash Paradigmalari (LO2)", "12"),
        ("    2.1 Uchta Paradigma Tushuntirildi", "12"),
        ("    2.2 Protsedural Dasturlash HotelOSda", "13"),
        ("    2.3 Ob'ektga Yo'naltirilgan Dasturlash HotelOSda", "14"),
        ("    2.4 Hodisaga Asoslangan Dasturlash HotelOSda", "17"),
        ("    2.5 Foydalanilgan Asosiy IDE Komponentlari", "19"),
        ("3-Vazifa — HotelOSni Qurish (LO3)", "20"),
        ("    3.1 Tizim Arxitekturasi va Komponentlari", "20"),
        ("    3.2 Xavfsizlik Mulohazalari", "22"),
        ("    3.3 IDE Ishlab Chiqish Jarayoni Dalili", "24"),
        ("    3.4 Test Stsenariylari va Chiqish Natijalari", "26"),
        ("4-Vazifa — Disk Raskadrovka va Kodlash Standartlari (LO4)", "28"),
        ("    4.1 Disk Raskadrovka Jarayoni", "28"),
        ("    4.2 Disk Raskadrovka Jurnali — Uchta Haqiqiy Xato", "29"),
        ("    4.3 Xavfsizlik uchun Disk Raskadrovka", "31"),
        ("    4.4 HotelOS Kodlash Standarti", "32"),
        ("    4.5 Kodlash Standartlari Jamoada Nima Uchun Muhim", "34"),
        ("Adabiyotlar Ro'yxati (Harvard formatida)", "36"),
        ("Qo'shimchalar", "37"),
    ]
    for title, page in items:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(16), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        run = p.add_run(title + "\t" + page)
        run.font.size = Pt(11)
