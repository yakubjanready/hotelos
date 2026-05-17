#!/usr/bin/env python3
"""HotelOS yakuniy Word hisobotini quradigan asosiy skript."""

from __future__ import annotations

import sys
from pathlib import Path

# Loyiha root'ini import yo'liga qo'shamiz
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document

from scripts.generate_report import (
    DOCS,
    add_cover_page,
    add_footer_page_numbers,
    add_heading,
    add_para,
    add_toc,
    page_break,
    setup_document_styles,
)
from scripts.report_task1 import task1
from scripts.report_task2 import task2
from scripts.report_task3 import task3
from scripts.report_task4 import task4
from scripts.report_references import references


def appendices(doc: Document) -> None:
    page_break(doc)
    add_heading(doc, "Qo'shimchalar", level=1)
    add_heading(doc, "A — Loyiha papka tuzilmasi", level=2)
    add_para(doc,
        "hotelos/\n"
        "├── README.md\n"
        "├── requirements.txt\n"
        "├── docker-compose.yml\n"
        "├── .env.example\n"
        "├── pyproject.toml\n"
        "├── shared/                      (umumiy modullar)\n"
        "│   ├── config.py\n"
        "│   ├── enums.py\n"
        "│   ├── events.py\n"
        "│   ├── models.py                (Pydantic modellari)\n"
        "│   ├── broker.py                (MessageBroker + Redis/InMemory)\n"
        "│   └── security.py              (InputValidator, JWT)\n"
        "├── reception_service/\n"
        "│   ├── algorithms.py            (assign_room, calculate_bill)\n"
        "│   └── main.py                  (FastAPI ilovasi)\n"
        "├── housekeeping_service/\n"
        "├── roomservice_service/\n"
        "├── maintenance_service/\n"
        "│   └── priority_queue.py        (heapq + FIFO)\n"
        "├── dashboard_service/\n"
        "│   ├── main.py                  (WebSocket + JWT)\n"
        "│   └── static/                  (HTML, JS, CSS)\n"
        "├── terminal_monitor/\n"
        "│   └── monitor.py               (rich TUI)\n"
        "├── data/\n"
        "│   ├── rooms.json               (120 xona inventari)\n"
        "│   └── menu.json\n"
        "├── tests/\n"
        "│   ├── test_algorithms.py       (14 ta birlik test)\n"
        "│   └── test_priority_queue.py   (5 ta test)\n"
        "├── scripts/\n"
        "│   ├── seed_rooms.py\n"
        "│   ├── demo_algorithms.py\n"
        "│   ├── run_test_scenarios.py    (TS-01..TS-08)\n"
        "│   ├── start_all.sh             (bir buyruq bilan ishga tushirish)\n"
        "│   └── stop_all.sh\n"
        "└── docs/\n"
        "    ├── diagram_*.svg / .png     (blok-sxemalar)\n"
        "    ├── screenshot_dashboard.png\n"
        "    └── test_results.md",
        indent=False, justify=False,
    )
    add_heading(doc, "B — Git majburiyat tarixi (kutilgan)", level=2)
    add_para(doc,
        "git log --oneline natijasi (10+ commit):\n\n"
        "feat: shared modules — config, enums, models, broker, security\n"
        "feat: room inventory seed data (120 rooms × 6 floors)\n"
        "feat: assign_room() multi-criteria algorithm with reasoning trace\n"
        "feat: calculate_bill() with Decimal precision + edge cases\n"
        "feat: reception service (check-in/check-out endpoints)\n"
        "feat: housekeeping service with FIFO cleaning queue\n"
        "feat: room service with order state machine\n"
        "feat: maintenance service with priority queue (heapq + FIFO)\n"
        "feat: dashboard service — WebSocket + JWT auth\n"
        "feat: dashboard frontend (HTML + Tailwind + vanilla JS)\n"
        "feat: terminal monitor (rich TUI)\n"
        "test: unit tests for algorithms + priority queue\n"
        "test: end-to-end scenarios TS-01..TS-08 runner\n"
        "fix: race condition — subscribe before start_listening (XATO-01)\n"
        "fix: priority queue FIFO tiebreaker (XATO-02)\n"
        "fix: Decimal/float mixing in billing (XATO-03)\n"
        "docs: README and start/stop scripts",
        indent=False, justify=False,
    )


def main() -> None:
    doc = Document()
    setup_document_styles(doc)
    add_footer_page_numbers(doc)

    # Cover + TOC
    add_cover_page(doc)
    page_break(doc)
    add_toc(doc)

    # Vazifalar
    task1(doc)
    task2(doc)
    task3(doc)
    task4(doc)

    # Manbalar va qo'shimchalar
    references(doc)
    appendices(doc)

    output = DOCS / "Yoqubjon_HotelOS_Hisobot.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"✓ Hisobot yaratildi: {output}")
    print(f"  Hajmi: {output.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
