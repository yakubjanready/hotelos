"""Adabiyotlar — Harvard referencing style."""

from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from scripts.generate_report import add_heading, page_break


def references(doc):
    page_break(doc)
    add_heading(doc, "Adabiyotlar Ro'yxati", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run(
        "Quyidagi manbalar Harvard referencing style (Cite Them Right, 11-nashr) "
        "qoidalariga muvofiq formatlangan. Manbalar mualliflarining familiyalari "
        "bo'yicha alfavit tartibida keltirilgan."
    ).font.size = Pt(11)
    doc.add_paragraph()

    refs = [
        # Books
        ("Aho, A. V., Hopcroft, J. E. and Ullman, J. D. (1987) "
         "*Data structures and algorithms*. Reading, MA: Addison-Wesley."),

        ("Fowler, M. (2018) *Refactoring: improving the design of existing code*. "
         "2nd edn. Boston, MA: Addison-Wesley."),

        ("Hunt, A. and Thomas, D. (2000) *The pragmatic programmer: from "
         "journeyman to master*. Reading, MA: Addison-Wesley."),

        ("McConnell, S. (2004) *Code complete: a practical handbook of software "
         "construction*. 2nd edn. Redmond, WA: Microsoft Press."),

        ("Newman, S. (2021) *Building microservices: designing fine-grained "
         "systems*. 2nd edn. Sebastopol, CA: O'Reilly Media."),

        # Online resources
        ("Mozilla Developer Network (2024) *WebSocket API*. Available at: "
         "https://developer.mozilla.org/en-US/docs/Web/API/WebSockets_API "
         "(Accessed: 12 May 2026)."),

        ("Python Software Foundation (2024) *The Python language reference, "
         "release 3.11*. Available at: https://docs.python.org/3.11/reference/ "
         "(Accessed: 10 May 2026)."),

        ("Redis Ltd. (2024) *Redis Pub/Sub documentation*. Available at: "
         "https://redis.io/docs/manual/pubsub (Accessed: 14 May 2026)."),

        ("Tiangolo, S. (2024) *FastAPI documentation*. Available at: "
         "https://fastapi.tiangolo.com (Accessed: 13 May 2026)."),

        ("van Rossum, G., Warsaw, B. and Coghlan, N. (2001) *PEP 8 — style "
         "guide for Python code*. Available at: "
         "https://peps.python.org/pep-0008/ (Accessed: 11 May 2026)."),

        ("Wiggins, A. (2017) *The Twelve-Factor App*. Available at: "
         "https://12factor.net (Accessed: 11 May 2026)."),
    ]

    for ref in refs:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.left_indent = Cm(1.0)
        para.paragraph_format.first_line_indent = Cm(-1.0)  # hanging indent
        para.paragraph_format.space_after = Pt(6)
        # parse *italic* markers
        parts = ref.split("*")
        for i, segment in enumerate(parts):
            run = para.add_run(segment)
            run.font.size = Pt(11)
            if i % 2 == 1:
                run.italic = True
